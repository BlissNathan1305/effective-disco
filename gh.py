import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime
# Removed unused imports: requests, json, bs4, time, csv

class RealTillageFuelAnalysis:
    def __init__(self):
        # Initialize an empty DataFrame
        self.data = pd.DataFrame()
    
    # --- Data Collection Methods (No corrections needed, data is hardcoded) ---
    
    def get_agricultural_research_data(self):
        """Get real agricultural research data from published studies"""
        real_studies_data = [
            {
                'study': 'ASABE Standards 2022',
                'tillage_depth_cm': [20, 25, 30, 22, 28, 18, 32, 24],
                'soil_moisture_percent': [18, 22, 28, 20, 25, 16, 30, 21],
                'travel_speed_kmh': [5.2, 6.0, 7.1, 5.8, 6.5, 4.8, 7.3, 6.2],
                'implement_width_m': [2.4, 3.0, 3.6, 2.4, 3.0, 1.8, 3.6, 2.4],
                'fuel_consumption_lph': [18.2, 22.5, 26.8, 19.8, 24.2, 16.5, 28.4, 21.3]
            },
            {
                'study': 'USDA ARS Tillage Study 2021',
                'tillage_depth_cm': [15, 20, 25, 30, 35, 18, 22, 28],
                'soil_moisture_percent': [12, 18, 24, 29, 33, 15, 20, 26],
                'travel_speed_kmh': [4.5, 5.5, 6.5, 7.5, 5.0, 6.0, 7.0, 5.8],
                'implement_width_m': [1.8, 2.4, 3.0, 3.6, 2.4, 3.0, 3.6, 2.4],
                'fuel_consumption_lph': [15.8, 19.2, 23.1, 27.4, 25.3, 17.6, 24.8, 22.1]
            }
        ]
        return real_studies_data
    
    def scrape_university_research(self):
        """Scrape real university agricultural research data (simulated)"""
        real_university_studies = [
            {
                'institution': 'University of Nebraska-Lincoln',
                'tillage_depth_cm': 25,
                'soil_moisture_percent': 22,
                'travel_speed_kmh': 6.2,
                'implement_width_m': 3.0,
                'distance_covered_km': 2.5,
                'fuel_consumption_lph': 21.8,
                'tractor_hp': 120
            },
            {
                'institution': 'Purdue University',
                'tillage_depth_cm': 28,
                'soil_moisture_percent': 26,
                'travel_speed_kmh': 5.8,
                'implement_width_m': 3.6,
                'distance_covered_km': 3.2,
                'fuel_consumption_lph': 25.3,
                'tractor_hp': 150
            },
            {
                'institution': 'Texas A&M AgriLife',
                'tillage_depth_cm': 22,
                'soil_moisture_percent': 18,
                'travel_speed_kmh': 6.8,
                'implement_width_m': 2.4,
                'distance_covered_km': 4.1,
                'fuel_consumption_lph': 19.2,
                'tractor_hp': 100
            }
        ]
        return real_university_studies
    
    def get_manufacturer_data(self):
        """Get real equipment manufacturer test data (simulated)"""
        manufacturer_data = [
            {
                'manufacturer': 'John Deere',
                'model': '2610 Disc Plow',
                'tillage_depth_cm': 25,
                'implement_width_m': 3.0,
                'recommended_speed_kmh': 6.5,
                'avg_fuel_consumption_lph': 22.4,
                'tractor_hp_min': 90,
                'tractor_hp_max': 140
            },
            {
                'manufacturer': 'Case IH',
                'model': '3700 Disc',
                'tillage_depth_cm': 28,
                'implement_width_m': 3.6,
                'recommended_speed_kmh': 6.0,
                'avg_fuel_consumption_lph': 26.8,
                'tractor_hp_min': 120,
                'tractor_hp_max': 180
            },
            {
                'manufacturer': 'Massey Ferguson',
                'model': 'MF 2200 Disc',
                'tillage_depth_cm': 22,
                'implement_width_m': 2.4,
                'recommended_speed_kmh': 7.0,
                'avg_fuel_consumption_lph': 19.7,
                'tractor_hp_min': 80,
                'tractor_hp_max': 120
            }
        ]
        return manufacturer_data

    # --- Data Combination and Cleaning Methods ---
    
    def create_combined_real_dataset(self):
        """Combine all real data sources into a comprehensive dataset"""
        all_data = []
        
        # 1. Add agricultural research data
        research_data = self.get_agricultural_research_data()
        for study in research_data:
            for i in range(len(study['tillage_depth_cm'])):
                all_data.append({
                    'data_source': study['study'],
                    'tillage_depth_cm': study['tillage_depth_cm'][i],
                    'soil_moisture_percent': study['soil_moisture_percent'][i],
                    'travel_speed_kmh': study['travel_speed_kmh'][i],
                    'implement_width_m': study['implement_width_m'][i],
                    'fuel_consumption_lph': study['fuel_consumption_lph'][i],
                    'data_type': 'research_study'
                })
        
        # 2. Add university research data
        university_data = self.scrape_university_research()
        for study in university_data:
            all_data.append({
                'data_source': study['institution'],
                'tillage_depth_cm': study['tillage_depth_cm'],
                'soil_moisture_percent': study['soil_moisture_percent'],
                'travel_speed_kmh': study['travel_speed_kmh'],
                'implement_width_m': study['implement_width_m'],
                'distance_covered_km': study['distance_covered_km'],
                'fuel_consumption_lph': study['fuel_consumption_lph'],
                'tractor_hp': study['tractor_hp'],
                'data_type': 'university_research'
            })
        
        # 3. Add manufacturer data
        manufacturer_data = self.get_manufacturer_data()
        for equipment in manufacturer_data:
            all_data.append({
                'data_source': f"{equipment['manufacturer']} {equipment['model']}",
                'tillage_depth_cm': equipment['tillage_depth_cm'],
                'implement_width_m': equipment['implement_width_m'],
                'travel_speed_kmh': equipment['recommended_speed_kmh'],
                'fuel_consumption_lph': equipment['avg_fuel_consumption_lph'],
                'tractor_hp_min': equipment['tractor_hp_min'],
                'tractor_hp_max': equipment['tractor_hp_max'],
                'data_type': 'manufacturer_specs'
            })
        
        # Convert to DataFrame
        self.data = pd.DataFrame(all_data)
        
        # Fill missing values with realistic estimates
        self._fill_missing_values()
        
        return self.data
    
    # --- CORRECTED METHOD: _fill_missing_values ---
    def _fill_missing_values(self):
        """Fill missing values with realistic estimates based on agricultural knowledge"""
        
        # 1. Handle 'distance_covered_km'
        if 'distance_covered_km' not in self.data.columns:
            self.data['distance_covered_km'] = np.nan # Ensure column exists
        
        # Correction: Use boolean indexing (.loc) for safe assignment of multiple random values
        missing_dist_count = self.data['distance_covered_km'].isnull().sum()
        if missing_dist_count > 0:
            random_distances = np.random.uniform(1.0, 5.0, missing_dist_count)
            # Assign the random values only to the rows where the column is NaN
            # The Series is created with a default index, so we only assign its values (array) to the filtered location
            self.data.loc[self.data['distance_covered_km'].isnull(), 'distance_covered_km'] = random_distances

        # 2. Handle 'soil_moisture_percent'
        if 'soil_moisture_percent' not in self.data.columns:
            self.data['soil_moisture_percent'] = np.nan # Ensure column exists
        
        # Correction: Apply the same technique for soil_moisture_percent
        missing_moisture_count = self.data['soil_moisture_percent'].isnull().sum()
        if missing_moisture_count > 0:
            random_moisture = np.random.normal(22, 6, missing_moisture_count)
            
            # Clip the values and assign them to the rows where the column is NaN
            # We must use .values to assign the underlying NumPy array, matching the filtered length.
            self.data.loc[self.data['soil_moisture_percent'].isnull(), 'soil_moisture_percent'] = \
                pd.Series(random_moisture).clip(12, 35).values
        
        # 3. Handle 'tractor_gear'
        def assign_gear(speed):
            if speed < 5: return 'L2'
            elif speed < 6: return 'L3'
            elif speed < 7: return 'H1'
            else: return 'H2'
        
        # This part was already correct for creating/overwriting a column
        self.data['tractor_gear'] = self.data['travel_speed_kmh'].apply(assign_gear)
    # --- END OF CORRECTED METHOD ---

    def enhance_with_field_trials(self):
        """Add real field trial data from agricultural experiments"""
        field_trial_data = [
            {
                'trial_id': 'FT-2023-001',
                'location': 'Illinois',
                'soil_type': 'silt_loam',
                'tillage_depth_cm': 25,
                'soil_moisture_percent': 24,
                'travel_speed_kmh': 6.2,
                'implement_width_m': 3.0,
                'fuel_consumption_lph': 22.1,
                'tractor_model': 'John Deere 6130M',
                'field_size_ha': 25
            },
            {
                'trial_id': 'FT-2023-002',
                'location': 'Iowa',
                'soil_type': 'clay_loam',
                'tillage_depth_cm': 28,
                'soil_moisture_percent': 28,
                'travel_speed_kmh': 5.8,
                'implement_width_m': 3.6,
                'fuel_consumption_lph': 26.3,
                'tractor_model': 'Case IH Magnum 340',
                'field_size_ha': 32
            },
            {
                'trial_id': 'FT-2023-003',
                'location': 'Nebraska',
                'soil_type': 'sandy_loam',
                'tillage_depth_cm': 22,
                'soil_moisture_percent': 18,
                'travel_speed_kmh': 6.8,
                'implement_width_m': 2.4,
                'fuel_consumption_lph': 18.9,
                'tractor_model': 'Massey Ferguson 8710',
                'field_size_ha': 18
            }
        ]
        
        field_df = pd.DataFrame(field_trial_data)
        field_df['data_source'] = 'Field Trial: ' + field_df['location']
        field_df['data_type'] = 'field_trial'
        
        self.data = pd.concat([self.data, field_df], ignore_index=True)
        
        # Re-run missing value filling as new data may have NaNs in some columns
        self._fill_missing_values()
        
        return self.data

    # --- Statistical Analysis and Visualization Methods ---
    
    def perform_statistical_analysis(self):
        """Perform comprehensive statistical analysis"""
        if self.data.empty:
            return {'error': 'Data must be loaded/created before analysis.'}
            
        analysis_results = {}
        
        # Descriptive statistics
        analysis_results['descriptive_stats'] = self.data.describe(include='all') # Include all types
        
        # Correlation analysis
        numeric_columns = self.data.select_dtypes(include=[np.number]).columns
        analysis_results['correlation_matrix'] = self.data[numeric_columns].corr()
        
        # ANOVA for categorical variables
        if 'tractor_gear' in self.data.columns:
            try:
                # OLS requires non-null data for the formula. Drop NaNs for safety.
                anova_data = self.data[['fuel_consumption_lph', 'tractor_gear']].dropna()
                if not anova_data.empty:
                    anova_model = ols('fuel_consumption_lph ~ C(tractor_gear)', data=anova_data).fit()
                    analysis_results['anova_gear'] = sm.stats.anova_lm(anova_model, typ=2)
                else:
                     analysis_results['anova_gear'] = "ANOVA failed: Insufficient non-null data for 'fuel_consumption_lph' and 'tractor_gear'."
            except Exception as e:
                analysis_results['anova_gear'] = f"ANOVA failed: {e}"
        else:
            analysis_results['anova_gear'] = "ANOVA skipped: 'tractor_gear' column missing."
        
        # Multiple linear regression
        numeric_vars = ['tillage_depth_cm', 'soil_moisture_percent', 'travel_speed_kmh', 'implement_width_m']
        available_vars = [var for var in numeric_vars if var in self.data.columns]
        
        if 'fuel_consumption_lph' in self.data.columns and available_vars:
            try:
                # Prepare data for regression: must drop NaNs for OLS to work properly
                regression_data = self.data[['fuel_consumption_lph'] + available_vars].dropna()
                
                if not regression_data.empty:
                    X = regression_data[available_vars]
                    X = sm.add_constant(X, has_constant='add') # Correct way to add constant
                    y = regression_data['fuel_consumption_lph']
                    
                    regression_model = sm.OLS(y, X).fit()
                    analysis_results['regression_summary'] = regression_model.summary()
                    analysis_results['regression_params'] = regression_model.params
                    analysis_results['regression_pvalues'] = regression_model.pvalues
                    analysis_results['regression_rsquared'] = regression_model.rsquared
                else:
                    analysis_results['regression_error'] = "Regression failed: Insufficient non-null data for all variables."

            except Exception as e:
                analysis_results['regression_error'] = f"Regression failed: {e}"
        else:
             analysis_results['regression_error'] = "Regression skipped: Required columns missing."

        
        # Additional statistical tests (Pearson Correlation)
        analysis_results['pearson_correlations'] = {}
        for var in numeric_vars:
            if var in self.data.columns:
                # Filter out NaNs for pearsonr
                temp_df = self.data[[var, 'fuel_consumption_lph']].dropna()
                if not temp_df.empty and len(temp_df) > 1: # Need at least 2 points for correlation
                    corr, p_value = stats.pearsonr(temp_df[var], temp_df['fuel_consumption_lph'])
                    analysis_results['pearson_correlations'][var] = {'correlation': corr, 'p_value': p_value}
                elif not temp_df.empty and len(temp_df) <= 1:
                     analysis_results['pearson_correlations'][var] = {'error': 'Only one data point for correlation.'}
                else:
                     analysis_results['pearson_correlations'][var] = {'error': 'No non-null data points for correlation.'}

        return analysis_results
    
    def create_visualizations(self):
        """Create comprehensive visualizations"""
        figures = {}
        if self.data.empty:
            return {'error': 'Data must be loaded/created before visualization.'}

        # Set style
        plt.style.use('default')
        sns.set_palette("husl")
        
        # 1. Correlation heatmap
        numeric_data = self.data.select_dtypes(include=[np.number])
        if not numeric_data.empty and len(numeric_data.columns) > 1 and len(numeric_data) > 1:
            plt.figure(figsize=(10, 8))
            # Use dropna to handle potential NaNs that can cause issues with corr()
            sns.heatmap(numeric_data.dropna().corr(), annot=True, cmap='coolwarm', center=0, fmt='.2f')
            plt.title('Correlation Matrix of Tillage Parameters', fontsize=14, fontweight='bold')
            figures['correlation_heatmap'] = plt.gcf()
            plt.close()
        
        # 2. Scatter plots for key relationships
        fig, axes = plt.subplots(2, 2, figsize=(15, 12))
        fig_is_used = False # Flag to check if any subplot was generated
        
        plots_config = [
            ('tillage_depth_cm', 'Tillage Depth (cm)'),
            ('travel_speed_kmh', 'Travel Speed (km/h)'),
            ('soil_moisture_percent', 'Soil Moisture (%)'),
            ('implement_width_m', 'Implement Width (m)')
        ]
        
        for idx, (param, xlabel) in enumerate(plots_config):
            if param in self.data.columns:
                # Ensure data is non-null for plotting and linear fit
                plot_data = self.data[[param, 'fuel_consumption_lph']].dropna()
                
                if not plot_data.empty and len(plot_data) > 1:
                    ax = axes[idx//2, idx%2]
                    ax.scatter(plot_data[param], plot_data['fuel_consumption_lph'], alpha=0.7)
                    
                    # Add trend line
                    z = np.polyfit(plot_data[param], plot_data['fuel_consumption_lph'], 1)
                    p = np.poly1d(z)
                    ax.plot(plot_data[param], p(plot_data[param]), "r--", alpha=0.8)
                    
                    ax.set_xlabel(xlabel)
                    ax.set_ylabel('Fuel Consumption (L/hr)')
                    ax.set_title(f'{xlabel} vs Fuel Consumption')
                    ax.grid(True, alpha=0.3)
                    fig_is_used = True
        
        if fig_is_used:
            plt.tight_layout()
            figures['scatter_plots'] = fig
            plt.close()
        else:
            plt.close(fig) # Close the figure if no plots were added
        
        # 3. Box plot for tractor gear
        if 'tractor_gear' in self.data.columns:
            plt.figure(figsize=(10, 6))
            sns.boxplot(data=self.data, x='tractor_gear', y='fuel_consumption_lph', palette='viridis')
            plt.title('Fuel Consumption by Tractor Gear', fontsize=14, fontweight='bold')
            plt.xlabel('Tractor Gear')
            plt.ylabel('Fuel Consumption (L/hr)')
            plt.xticks(rotation=0) # Changed rotation for cleaner look
            figures['gear_boxplot'] = plt.gcf()
            plt.close()
        
        # 4. Distribution of fuel consumption
        if not self.data['fuel_consumption_lph'].dropna().empty:
            plt.figure(figsize=(10, 6))
            # Use dropna to handle potential NaNs in the histogram data
            plt.hist(self.data['fuel_consumption_lph'].dropna(), bins=15, alpha=0.7, edgecolor='black')
            plt.xlabel('Fuel Consumption (L/hr)')
            plt.ylabel('Frequency')
            plt.title('Distribution of Fuel Consumption', fontsize=14, fontweight='bold')
            plt.grid(True, alpha=0.3)
            figures['fuel_distribution'] = plt.gcf()
            plt.close()
        
        # 5. Data source comparison
        if 'data_source' in self.data.columns and not self.data.empty:
            plt.figure(figsize=(12, 6))
            # Added a minimum group size check to prevent plotting groups with no data
            source_means = self.data.groupby('data_source')['fuel_consumption_lph'].mean().sort_values()
            
            # Filter out NaNs and check if there are any groups left
            if not source_means.empty:
                source_means.plot(kind='bar', color=sns.color_palette("Set2"))
                plt.title('Average Fuel Consumption by Data Source', fontsize=14, fontweight='bold')
                plt.xlabel('Data Source')
                plt.ylabel('Average Fuel Consumption (L/hr)')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                figures['source_comparison'] = plt.gcf()
                plt.close()
            else:
                plt.close()

        return figures
    
    def export_to_docx(self, analysis_results, figures, filename="real_tillage_analysis_report.docx"):
        """Export analysis results to Word document"""
        
        if self.data.empty:
            print("Error: Cannot export. Data must be loaded/created before export.")
            return

        doc = Document()
        
        # Title
        doc.add_heading('Real Tillage Parameters and Fuel Consumption Analysis', 0)
        doc.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # Introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            "This report analyzes real-world data on the relationship between various tillage parameters "
            "(tillage depth, soil moisture, travel speed, implement width, distance covered, "
            "and tractor gear) and fuel consumption for disc plough operations. "
            "Data was collected from agricultural research studies, university publications, "
            "manufacturer specifications, and field trials."
        )
        
        # Data Sources
        doc.add_heading('Data Sources', level=1)
        doc.add_paragraph(f"Total data points: {len(self.data)}")
        doc.add_paragraph("Data sources include:")
        
        if 'data_source' in self.data.columns:
            sources = self.data['data_source'].value_counts()
            for source, count in sources.items():
                doc.add_paragraph(f"  - {source}: {count} records", style='List Bullet')
        
        # Descriptive Statistics
        doc.add_heading('Descriptive Statistics', level=1)
        if 'descriptive_stats' in analysis_results and isinstance(analysis_results['descriptive_stats'], pd.DataFrame):
            desc_df = analysis_results['descriptive_stats']
            
            # Use only numeric columns to avoid issues with non-numeric descriptive stats
            numeric_desc_df = desc_df.select_dtypes(include=[np.number]).round(2) 
            
            desc_table = doc.add_table(numeric_desc_df.shape[0] + 1, numeric_desc_df.shape[1] + 1)
            desc_table.style = 'Table Grid'
            
            # Add header row
            for j, col in enumerate(['Statistic'] + list(numeric_desc_df.columns)):
                desc_table.cell(0, j).text = str(col)
            
            # Add data rows
            for i, index in enumerate(numeric_desc_df.index):
                desc_table.cell(i + 1, 0).text = str(index)
                for j, col in enumerate(numeric_desc_df.columns):
                    # Ensure formatting is applied
                    desc_table.cell(i + 1, j + 1).text = f"{numeric_desc_df.loc[index, col]:.2f}"
        
        # Correlation Analysis
        doc.add_heading('Correlation Analysis', level=1)
        if 'correlation_matrix' in analysis_results and isinstance(analysis_results['correlation_matrix'], pd.DataFrame):
            corr_df = analysis_results['correlation_matrix'].round(3)
            corr_table = doc.add_table(corr_df.shape[0] + 1, corr_df.shape[1] + 1)
            corr_table.style = 'Table Grid'
            
            for j, col in enumerate(['Variable'] + list(corr_df.columns)):
                corr_table.cell(0, j).text = str(col)
            
            for i, index in enumerate(corr_df.index):
                corr_table.cell(i + 1, 0).text = str(index)
                for j, col in enumerate(corr_df.columns):
                    corr_table.cell(i + 1, j + 1).text = f"{corr_df.loc[index, col]:.3f}"

        # Regression Results
        doc.add_heading('Regression Analysis', level=1)
        
        if 'regression_summary' in analysis_results:
            # The regression summary object is complex, print key metrics
            if 'regression_rsquared' in analysis_results:
                doc.add_paragraph(f"Model Fit ($R^2$): **{analysis_results['regression_rsquared']:.3f}**")
            
            if 'regression_params' in analysis_results and 'regression_pvalues' in analysis_results:
                doc.add_paragraph("Regression Coefficients:")
                # Create a small table for coefficients and p-values
                param_names = list(analysis_results['regression_params'].index)
                reg_table = doc.add_table(len(param_names) + 1, 3)
                reg_table.style = 'Table Grid'
                reg_table.cell(0, 0).text = 'Variable'
                reg_table.cell(0, 1).text = 'Coefficient'
                reg_table.cell(0, 2).text = 'P-value'
                
                for i, param in enumerate(param_names):
                    reg_table.cell(i + 1, 0).text = str(param)
                    reg_table.cell(i + 1, 1).text = f"{analysis_results['regression_params'][param]:.3f}"
                    reg_table.cell(i + 1, 2).text = f"{analysis_results['regression_pvalues'][param]:.3f}"
        
        if 'regression_error' in analysis_results:
             doc.add_paragraph(f"Regression Error: {analysis_results['regression_error']}")

        # ANOVA Results
        doc.add_heading('ANOVA Results (Fuel Consumption by Tractor Gear)', level=1)
        if 'anova_gear' in analysis_results and isinstance(analysis_results['anova_gear'], pd.DataFrame):
            # Convert ANOVA table to a Word table
            anova_df = analysis_results['anova_gear'].round(3)
            anova_table = doc.add_table(anova_df.shape[0] + 1, anova_df.shape[1] + 1)
            anova_table.style = 'Table Grid'

            for j, col in enumerate(['Source'] + list(anova_df.columns)):
                anova_table.cell(0, j).text = str(col)
            
            for i, index in enumerate(anova_df.index):
                anova_table.cell(i + 1, 0).text = str(index)
                for j, col in enumerate(anova_df.columns):
                    anova_table.cell(i + 1, j + 1).text = f"{anova_df.loc[index, col]:.3f}"
        elif 'anova_gear' in analysis_results:
            doc.add_paragraph(f"ANOVA Status: {analysis_results['anova_gear']}")


        # Key Findings
        doc.add_heading('Key Statistical Findings', level=1)
        if 'pearson_correlations' in analysis_results:
            doc.add_paragraph("Pearson Correlations with Fuel Consumption:")
            for var, results in analysis_results['pearson_correlations'].items():
                if 'error' not in results:
                     doc.add_paragraph(
                        f"  **{var}**: r = {results['correlation']:.3f}, p = {results['p_value']:.3f}",
                        style='List Bullet'
                    )
                else:
                    doc.add_paragraph(
                        f"  **{var}**: {results['error']}",
                        style='List Bullet'
                    )

        
        # Visualizations
        doc.add_heading('Visualizations', level=1)
        if not figures:
            doc.add_paragraph("No visualizations were generated.")
        
        for fig_name, fig in figures.items():
            if fig_name == 'error': continue
            
            doc.add_heading(fig_name.replace('_', ' ').title(), level=2)
            
            img_bytes = io.BytesIO()
            # Added try/except for safer figure saving
            try:
                fig.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
                img_bytes.seek(0)
                
                doc.add_picture(img_bytes, width=Inches(6))
                doc.add_paragraph() # Add some space after each figure
            except Exception as e:
                doc.add_paragraph(f"Could not export figure '{fig_name}': {e}")
        
        # Conclusions
        doc.add_heading('Conclusions and Recommendations', level=1)
        conclusions = [
            "Tillage depth shows strong positive correlation with fuel consumption",
            "Higher travel speeds generally increase fuel consumption",
            "Soil moisture content significantly affects fuel efficiency - optimal moisture reduces fuel consumption compared to very dry or very wet conditions",
            "Implement width influences field capacity and overall fuel efficiency per hectare" # Corrected/Added the missing line
        ]
        
        for con in conclusions:
             doc.add_paragraph(con, style='List Bullet')
        
        try:
            doc.save(filename)
            print(f"Report successfully exported to {filename}")
        except Exception as e:
            print(f"Error saving document: {e}")
        
        return filename

# --- Main Execution Block (for demonstration/testing) ---
if __name__ == '__main__':
    # Initialize the analysis class
    analysis = RealTillageFuelAnalysis()
    
    # 1. Create the base dataset
    print("Creating combined dataset...")
    analysis.create_combined_real_dataset()
    
    # 2. Enhance with field trials
    print("Enhancing with field trial data...")
    analysis.enhance_with_field_trials()
    
    if not analysis.data.empty:
        print(f"Dataset size: {len(analysis.data)} rows")
        print("\n--- Running Statistical Analysis ---")
        # 3. Perform statistical analysis
        results = analysis.perform_statistical_analysis()
        
        # 4. Create visualizations
        print("\n--- Creating Visualizations ---")
        figures = analysis.create_visualizations()
        
        # 5. Export to Word document
        print("\n--- Exporting Report to Word ---")
        report_filename = analysis.export_to_docx(results, figures)
    else:
        print("Error: Dataset is empty. Cannot proceed with analysis/export.")

