import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from io import BytesIO
from scipy.stats import linregress
from datetime import datetime

# --- Configuration and Setup ---
FILE_NAME = 'youtube-top-100-songs-2025.csv'
REPORT_OUTPUT = 'YouTube_Top_Songs_2025_Report.docx'

# Set up plotting style
sns.set_style("whitegrid")

# --- Helper Function for DOCX Tables ---
def add_df_to_document(document, df, title):
    """Adds a DataFrame as a formatted table to the Word document."""
    document.add_heading(title, level=2)
    
    # Reset index to include it as a column if it has meaningful names (like 'channel')
    df_temp = df.reset_index()
    
    table = document.add_table(df_temp.shape[0] + 1, df_temp.shape[1], style='Table Grid')
    
    # Add column headers
    for j, col in enumerate(df_temp.columns):
        table.cell(0, j).text = str(col)

    # Add data rows
    for i, row in df_temp.iterrows():
        for j, cell in enumerate(row):
            # Format large numbers for readability in the document
            if isinstance(cell, (int, float)) and cell > 10000:
                # Format to a readable string (e.g., 2.0B, 159.6M)
                if cell >= 1e9:
                    formatted_cell = f"{cell/1e9:.2f}B"
                elif cell >= 1e6:
                    formatted_cell = f"{cell/1e6:.1f}M"
                elif cell >= 1e3:
                    formatted_cell = f"{cell/1e3:.0f}K"
                else:
                    formatted_cell = str(cell)
            else:
                formatted_cell = str(cell)
            
            table.cell(i + 1, j).text = formatted_cell

    document.add_paragraph() # Add a space after the table

# --- Main Analysis Script ---
def run_analysis_and_report():
    print(f"--- Starting Analysis of {FILE_NAME} ---")
    
    # 1. Data Loading and Cleaning
    try:
        df = pd.read_csv(FILE_NAME)
    except FileNotFoundError:
        print(f"Error: The file '{FILE_NAME}' was not found. Please ensure it is in the same directory.")
        return

    print("Data loaded successfully.")
    
    # Convert and clean numeric columns
    numeric_cols = ['view_count', 'channel_follower_count', 'duration']
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors='coerce')
        
    df.dropna(subset=numeric_cols, inplace=True)
    df['view_count'] = df['view_count'].astype(np.int64)
    df['channel_follower_count'] = df['channel_follower_count'].astype(np.int64)

    print(f"Data cleaned. {len(df)} rows remaining for analysis.")

    # 2. Descriptive Statistics and Correlation
    descriptive_stats = df[['view_count', 'duration', 'channel_follower_count']].describe().T
    
    correlation = df[['view_count', 'channel_follower_count']].corr().iloc[0, 1]
    
    # 3. Categorical Analysis (Tags)
    df['tags'] = df['tags'].fillna('')
    all_tags = df['tags'].str.split(';').explode()
    all_tags = all_tags.str.strip()
    all_tags = all_tags[all_tags != '']
    top_tags = all_tags.value_counts().nlargest(10).reset_index()
    top_tags.columns = ['Tag', 'Count']

    # 4. Channel Performance Analysis
    channel_performance = df.groupby('channel').agg(
        Total_Views=('view_count', 'sum'),
        Total_Songs=('title', 'count'),
        Avg_Views_per_Song=('view_count', 'mean'),
        Followers=('channel_follower_count', 'first')
    ).sort_values(by='Total_Views', ascending=False)
    
    top_10_channels = channel_performance.nlargest(10, 'Total_Views')

    # 5. Plotting and Saving Figures
    
    # List to store plot paths for DOCX
    plot_paths = []

    # --- Plot 1: View Count Distribution (Log Scale) ---
    plt.figure(figsize=(10, 6))
    sns.histplot(df['view_count'], bins=20, kde=True, log_scale=True, color='skyblue')
    plt.title('Distribution of Song View Counts (Log Scale)', fontsize=16)
    plt.xlabel('View Count (Log Scale)', fontsize=12)
    plt.ylabel('Frequency', fontsize=12)
    plt.tight_layout()
    plt.savefig('view_count_log_distribution.png')
    plot_paths.append(('view_count_log_distribution.png', "Figure 1: Distribution of View Counts"))
    plt.close()
    
    # --- Plot 2: Scatter Plot (View Count vs. Channel Followers) ---
    plt.figure(figsize=(10, 6))
    sns.regplot(x='channel_follower_count', y='view_count', data=df, 
                scatter_kws={'alpha':0.6, 'color':'darkblue'}, line_kws={'color':'red', 'linestyle':'--'})
    plt.xscale('log')
    plt.yscale('log')
    plt.title('Relationship: Channel Followers vs. Song View Count (Log-Log)', fontsize=16)
    plt.xlabel('Channel Follower Count (Log Scale)', fontsize=12)
    plt.ylabel('View Count (Log Scale)', fontsize=12)
    plt.tight_layout()
    plt.savefig('follower_vs_view_scatter.png')
    plot_paths.append(('follower_vs_view_scatter.png', "Figure 2: View Count vs. Channel Followers"))
    plt.close()

    # --- Plot 3: Top 10 Most Frequent Tags ---
    plt.figure(figsize=(12, 6))
    sns.barplot(x=top_tags['Tag'], y=top_tags['Count'], palette="viridis")
    plt.title('Top 10 Most Frequent Tags in Top 100 Songs', fontsize=16)
    plt.xlabel('Tag', fontsize=12)
    plt.ylabel('Count', fontsize=12)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('top_tags_bar_chart.png')
    plot_paths.append(('top_tags_bar_chart.png', "Figure 3: Top 10 Most Frequent Tags"))
    plt.close()

    # 6. DOCX Report Generation
    
    document = Document()
    document.add_heading('YouTube Top 100 Songs 2025: Statistical Analysis Report', 0)
    document.add_paragraph(f'Report Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    document.add_paragraph(f'Data Source: {FILE_NAME}')
    document.add_page_break()

    # --- Section 1: Descriptive Statistics ---
    document.add_heading('1. Descriptive Statistics', level=1)
    document.add_paragraph('Summary statistics for the primary quantitative metrics, highlighting the distribution skewness.')
    
    # Format descriptive stats for better table presentation
    stats_df = descriptive_stats.reset_index()
    stats_df.columns = ['Metric', 'Count', 'Mean', 'Std Dev', 'Min', '25%', 'Median', '75%', 'Max']
    stats_df['Metric'] = stats_df['Metric'].replace({
        'view_count': 'View Count',
        'duration': 'Duration (sec)',
        'channel_follower_count': 'Channel Followers'
    })
    
    add_df_to_document(document, stats_df, 'Table 1: Key Descriptive Statistics')
    
    document.add_paragraph(
        'The massive difference between the Mean View Count and the Median confirms the distribution is highly '
        'right-skewed, meaning a few exceptionally viral hits (outliers) heavily influence the overall average.'
    )

    # --- Section 2: Correlation Analysis ---
    document.add_heading('2. Correlation Analysis', level=1)
    
    document.add_paragraph(
        f'The Pearson correlation coefficient between **View Count** and **Channel Follower Count** is **{correlation:.3f}**.'
    )
    document.add_paragraph(
        'This indicates a very weak positive correlation. While channels with a larger follower base tend to have '
        'a baseline advantage, the view count of a song is only slightly related to the channel size. '
        'Viral factors, external marketing, and the inherent popularity of the song itself are likely '
        'more dominant drivers of ultra-high view counts.'
    )
    
    # --- Section 3: Channel Performance ---
    document.add_heading('3. Top Channel Performance', level=1)
    
    top_10_channels_report = top_10_channels[['Total_Views', 'Total_Songs', 'Avg_Views_per_Song', 'Followers']]
    add_df_to_document(document, top_10_channels_report, 'Table 2: Top 10 Channels by Total View Count')
    
    document.add_paragraph(
        'The analysis shows ROSÉ and Lady Gaga lead the chart in total views, but channels like Sabrina Carpenter '
        'demonstrate high consistency with the most songs in the Top 100. This table highlights channels that '
        'drive massive traffic through one or two super-hits versus those that consistently produce chart-worthy content.'
    )

    # --- Section 4: Categorical Insights (Tags) ---
    document.add_heading('4. Top Tags and Genre Insights', level=1)
    add_df_to_document(document, top_tags, 'Table 3: Top 10 Most Frequent YouTube Tags')

    document.add_paragraph(
        'The overwhelming dominance of the "Pop" tag confirms the genre preference of the current Top 100 list. '
        'The consistent appearance of record labels and specific artist/fandom names in the tags indicates '
        'effective use of keywords for search engine optimization (SEO) on YouTube.'
    )

    # --- Section 5: Visualizations ---
    document.add_heading('5. Visualizations', level=1)
    
    for path, caption in plot_paths:
        document.add_heading(caption, level=3)
        try:
            document.add_picture(path, width=Inches(6.5))
        except FileNotFoundError:
            document.add_paragraph(f"Could not find figure file: {path}")

    # --- Save Report ---
    document.save(REPORT_OUTPUT)
    print(f"\n--- Report '{REPORT_OUTPUT}' generated successfully! ---")

if __name__ == '__main__':
    run_analysis_and_report()


