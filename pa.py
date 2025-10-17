# pollution_analysis.py  (CORRECTED 2025-10-17)
# ------------------------------------------------------------
# Author : you
# Data   : McDonald & Schwing (1973) pollution & mortality
# Task   : extensive EDA + inferential stats → Word report
# ------------------------------------------------------------
import os, textwrap, warnings
import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import statsmodels.api as sm

from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.diagnostic import het_breuschpagan
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import RidgeCV, LassoCV
from docx import Document
from docx.shared import Inches

sns.set_theme(style="whitegrid", palette="colorblind")
plt.rcParams["figure.figsize"] = (6, 4)
warnings.filterwarnings("ignore", category=FutureWarning)

# ------------------------------------------------------------------
# 1. READ & PARSE THE RAW FILE  ––  FIXED-WIDTH VERSION
# ------------------------------------------------------------------
var_names = [
    "PREC","JANT","JULT","OVR65","POPN","EDUC","HOUS","DENS",
    "NONW","WWDRK","POOR","HC","NOX","SO2","HUMID","MORT"
]

colspecs = [(0,5), (5,11), (11,17), (17,24), (24,31), (31,37),
            (37,44), (44,51), (51,58), (58,64), (64,70), (70,76),
            (76,82), (82,88), (88,94), (94,101)]

df = pd.read_fwf("pollution.txt", colspecs=colspecs, header=None,
                 names=var_names, na_values=".")

# CRITICAL CORRECTION 1: Explicitly coerce all columns to numeric.
# This resolves the 'could not convert string to float: "This"' ValueError
# by replacing any remaining non-numeric text (like 'This') with NaN.
df = df.apply(pd.to_numeric, errors='coerce') 

df = df.dropna(how="all")          # drop any all-empty rows
print(f"Shape after parsing: {df.shape}")

# ------------------------------------------------------------------
# 2. QUICK DESCRIPTIVE TABLE
# ------------------------------------------------------------------
# This section now runs correctly because all columns are numeric
desc = df.describe().T
desc["missing"] = df.isna().sum()
desc["skew"]    = df.skew()
desc["kurt"]    = df.kurtosis()

# ------------------------------------------------------------------
# 3. GRAPH FUNCTIONS (No changes needed)
# ------------------------------------------------------------------
def save(name):
    plt.tight_layout()
    plt.savefig(name, dpi=300)
    plt.close()

# 3a. Pair-plot (sample down if too big)
sns.pairplot(df.sample(min(60, len(df))), corner=True)
save("pairplot.png")

# 3b. Correlation heat-map
plt.figure()
mask = np.triu(np.ones_like(df.corr(), dtype=bool))
sns.heatmap(df.corr(), mask=mask, annot=True, fmt=".2f", cmap="coolwarm")
save("corr_heatmap.png")

# 3c. Distribution of target
plt.figure()
sns.histplot(df["MORT"], kde=True)
plt.title("Distribution of MORT")
save("mort_dist.png")

# 3d. Box-plots for pollution variables
pollution_vars = ["HC","NOX","SO2"]
plt.figure()
sns.boxplot(data=df[pollution_vars])
plt.title("Box-plots of pollution variables")
save("pollution_box.png")

# ------------------------------------------------------------------
# 4. INFERENTIAL PART
# ------------------------------------------------------------------
y = df["MORT"]
X_base = df.drop(columns=["MORT"]) # Non-standardized predictors

# Standardize X
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X_base)
X_scaled = pd.DataFrame(X_scaled, columns=X_base.columns)

# Add constant for OLS and VIF (This is the full design matrix)
X_const = sm.add_constant(X_scaled) 

# 4a. OLS
# Use the full design matrix (X_const) for OLS
ols = sm.OLS(y, X_const).fit() 
ols_summary = ols.summary().as_text()

# 4b. VIF frame - CRITICAL CORRECTION 2: VIF must be calculated on the full design matrix (X_const)
# and the VIF for the constant must be removed.
vif_df = pd.DataFrame({
    "var" : X_const.columns,
    "VIF" : [variance_inflation_factor(X_const.values, i) for i in range(X_const.shape[1])]
})

# Drop VIF for the 'const' term (intercept)
vif_df = vif_df.drop(vif_df[vif_df['var'] == 'const'].index)


# 4c. Heteroscedasticity
bp_lm, bp_pval, _, _ = het_breuschpagan(ols.resid, ols.model.exog)

# 4d. Regularisation paths
# Use X_scaled (standardized, no constant) for sklearn models
alphas = np.logspace(-2, 3, 100)
ridge = RidgeCV(alphas=alphas, cv=5).fit(X_scaled, y)
lasso = LassoCV(alphas=alphas, cv=5).fit(X_scaled, y)

# plot regularisation path for Ridge
plt.figure()
coefs = []
for a in alphas:
    # Use X_scaled here
    coefs.append(RidgeCV(alphas=[a]).fit(X_scaled, y).coef_) 
ax = plt.gca()
ax.plot(alphas, coefs)
ax.set_xscale("log")
ax.set_xlabel("α")
ax.set_ylabel("Standardised coefficient")
ax.set_title("Ridge trace – pollution data")
save("ridge_trace.png")

# ------------------------------------------------------------------
# 5. WRITE WORD DOCUMENT (No changes needed)
# ------------------------------------------------------------------
doc = Document()
doc.add_heading("Pollution & Mortality – Statistical Report", 0)

doc.add_heading("1. Data overview", level=1)
doc.add_paragraph(f"The file contains {df.shape[0]} SMSA observations and {df.shape[1]} variables.")
doc.add_paragraph("Table 1 – Descriptive statistics")
t = doc.add_table(rows=1, cols=len(desc.columns)+1)
t.style = "Light List Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "Variable"
for j, col in enumerate(desc.columns):
    hdr[j+1].text = str(col)
for idx, row in desc.iterrows():
    cells = t.add_row().cells
    cells[0].text = idx
    for j, v in enumerate(row):
        cells[j+1].text = f"{v:.2f}"

def add_fig(path, caption):
    doc.add_picture(path, width=Inches(5))
    doc.add_paragraph(caption, style="Caption")

doc.add_heading("2. Exploratory graphics", level=1)
add_fig("pairplot.png", "Figure 1 – Scatter-plot matrix (corner)")
add_fig("corr_heatmap.png", "Figure 2 – Pearson correlation heat-map")
add_fig("mort_dist.png", "Figure 3 – Distribution of age-adjusted mortality")
add_fig("pollution_box.png", "Figure 4 – Pollution variables box-plots")

doc.add_heading("3. Regression diagnostics", level=1)
doc.add_paragraph("Table 2 – Variance inflation factors")
t = doc.add_table(rows=1, cols=2)
t.style = "Light List Accent 1"
t.rows[0].cells[0].text = "Variable"
t.rows[0].cells[1].text = "VIF"
for _, row in vif_df.iterrows():
    cells = t.add_row().cells
    cells[0].text = row["var"]
    cells[1].text = f"{row['VIF']:.2f}"

doc.add_paragraph(f"Breusch-Pagan test LM = {bp_lm:.2f}, p = {bp_pval:.3g}")
doc.add_paragraph("Table 3 – OLS summary (full text)")
doc.add_paragraph(textwrap.indent(ols_summary, prefix="    "))

add_fig("ridge_trace.png", "Figure 5 – Ridge trace plot")

doc.save("pollution_report.docx")
print("Report exported → pollution_report.docx")

