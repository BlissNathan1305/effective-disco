#!/usr/bin/env python3
"""
tractor_fuel_analysis.py

Generates synthetic data for tractor fuel consumption with a spring tine cultivator,
runs statistical analysis (OLS with interactions), saves diagnostic plots, and
exports a DOCX report with tables and images.

Outputs:
 - synthetic_tractor_fuel_data.csv
 - residuals_vs_fitted.png
 - qqplot_residuals.png
 - scatter_matrix.png
 - tractor_fuel_analysis_report.docx
"""

import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document
from docx.shared import Inches

# -------------------------
# Configuration / settings
# -------------------------
OUT_DIR = os.path.abspath("./output")
os.makedirs(OUT_DIR, exist_ok=True)

CSV_PATH = os.path.join(OUT_DIR, "synthetic_tractor_fuel_data.csv")
DOCX_PATH = os.path.join(OUT_DIR, "tractor_fuel_analysis_report.docx")
IMG_RESID = os.path.join(OUT_DIR, "residuals_vs_fitted.png")
IMG_QQ = os.path.join(OUT_DIR, "qqplot_residuals.png")
IMG_SCAT = os.path.join(OUT_DIR, "scatter_matrix.png")

RANDOM_SEED = 42
N_SAMPLES = 300

# -------------------------
# 1) Data generation
# -------------------------
def generate_synthetic_data(n=N_SAMPLES, seed=RANDOM_SEED):
    np.random.seed(seed)
    # realistic ranges
    tillage_depth = np.random.uniform(5, 25, n)           # cm
    implement_width = np.random.uniform(1.5, 4.0, n)      # m
    soil_moisture = np.random.uniform(10, 35, n)          # % volumetric
    speed = np.random.uniform(2, 8, n)                    # km/h

    # Underlying (plausible) model for fuel consumption (L/h)
    fuel_base = 8.0
    fuel = (
        fuel_base
        + 0.12 * tillage_depth
        + 0.9  * implement_width
        + 0.05 * soil_moisture
        + 0.6  * speed
        + 0.02 * tillage_depth * implement_width
        + 0.015 * soil_moisture * speed
    )

    # heteroscedastic noise: increases with depth and speed
    noise_sd = 0.8 + 0.02 * tillage_depth + 0.05 * speed
    noise = np.random.normal(0, noise_sd)
    # draw vector noise
    noise = np.random.normal(0, noise_sd, n)
    fuel_observed = fuel + noise

    df = pd.DataFrame({
        "tillage_depth_cm": np.round(tillage_depth, 2),
        "implement_width_m": np.round(implement_width, 2),
        "soil_moisture_pct": np.round(soil_moisture, 2),
        "speed_kmh": np.round(speed, 2),
        "fuel_L_per_hr": np.round(fuel_observed, 3),
    })
    return df

# -------------------------
# 2) Analysis functions
# -------------------------
def fit_model(df):
    # formula with two interactions
    formula = ("fuel_L_per_hr ~ tillage_depth_cm + implement_width_m + "
               "soil_moisture_pct + speed_kmh + "
               "tillage_depth_cm:implement_width_m + soil_moisture_pct:speed_kmh")
    model = smf.ols(formula=formula, data=df).fit()
    return model, formula

def compute_vif(df):
    # VIF for main effects (include constant)
    X = df[["tillage_depth_cm", "implement_width_m", "soil_moisture_pct", "speed_kmh"]].copy()
    X = sm.add_constant(X)
    vif = []
    for i in range(X.shape[1]):
        try:
            v = variance_inflation_factor(X.values, i)
        except Exception:
            v = np.nan
        vif.append(v)
    vif_df = pd.DataFrame({"variable": X.columns, "VIF": vif})
    return vif_df

# -------------------------
# 3) Plots
# -------------------------
def make_plots(model, df):
    residuals = model.resid
    fitted = model.fittedvalues

    # Residuals vs fitted
    plt.figure(figsize=(7,5))
    plt.scatter(fitted, residuals, alpha=0.6)
    plt.axhline(0, linestyle='--', linewidth=0.8)
    plt.xlabel("Fitted values (L/h)")
    plt.ylabel("Residuals (L/h)")
    plt.title("Residuals vs Fitted")
    plt.tight_layout()
    plt.savefig(IMG_RESID)
    plt.close()

    # QQ plot
    plt.figure(figsize=(6,6))
    sm.qqplot(residuals, line='45', fit=True)
    plt.title("QQ-plot of residuals")
    plt.tight_layout()
    plt.savefig(IMG_QQ)
    plt.close()

    # Scatter matrix (pairwise)
    # Use a subset of variables for readability
    pd.plotting.scatter_matrix(df[["tillage_depth_cm", "implement_width_m", "soil_moisture_pct", "speed_kmh", "fuel_L_per_hr"]],
                               diagonal='hist', figsize=(10,10))
    plt.suptitle("Scatter matrix - variables", y=0.92)
    plt.tight_layout()
    plt.savefig(IMG_SCAT)
    plt.close()

# -------------------------
# 4) DOCX report generation
# -------------------------
def create_docx_report(df, desc, corr, model, formula, anova_table, vif_df):
    doc = Document()
    doc.add_heading("Statistical Analysis Report: Tractor Fuel Consumption (Synthetic Data)", level=1)

    doc.add_paragraph(
        "This report contains a synthetic dataset generated to study the impact of "
        "tillage depth, implement width, soil moisture, and speed on tractor fuel consumption (L/h) "
        "for a spring tine cultivator. The dataset contains {} simulated observations.".format(len(df))
    )

    doc.add_heading("1. Data preview", level=2)
    doc.add_paragraph("First 8 rows of the dataset:")
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.head(8).iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])

    doc.add_heading("2. Descriptive statistics", level=2)
    for idx, row in desc.iterrows():
        doc.add_paragraph(f"{idx}: mean={row['mean']:.3f}, std={row['std']:.3f}, min={row['min']:.3f}, 25%={row['25%']:.3f}, 50%={row['50%']:.3f}, 75%={row['75%']:.3f}, max={row['max']:.3f}")

    doc.add_heading("3. Correlation matrix", level=2)
    doc.add_paragraph("Pearson correlation between variables:")
    corr_table = doc.add_table(rows=1, cols=len(corr.columns)+1)
    corr_table.rows[0].cells[0].text = ""
    for i, col in enumerate(corr.columns):
        corr_table.rows[0].cells[i+1].text = col
    for i, idx in enumerate(corr.index):
        row_cells = corr_table.add_row().cells
        row_cells[0].text = idx
        for j, col in enumerate(corr.columns):
            row_cells[j+1].text = f"{corr.iloc[i,j]:.3f}"

    doc.add_heading("4. Regression model", level=2)
    doc.add_paragraph(f"Model formula: {formula}")
    doc.add_paragraph("Model summary (coefficients):")
    reg_table = doc.add_table(rows=1, cols=6)
    reg_table.rows[0].cells[0].text = "term"
    reg_table.rows[0].cells[1].text = "coef"
    reg_table.rows[0].cells[2].text = "std err"
    reg_table.rows[0].cells[3].text = "t"
    reg_table.rows[0].cells[4].text = "P>|t|"
    reg_table.rows[0].cells[5].text = "[0.025 0.975]"
    sum2 = model.summary2().tables[1]
    for term, row in sum2.iterrows():
        r = reg_table.add_row().cells
        r[0].text = str(term)
        r[1].text = f"{row['Coef.']:.4f}"
        r[2].text = f"{row['Std.Err.']:.4f}"
        r[3].text = f"{row['t']:.3f}"
        r[4].text = f"{row['P>|t|']:.4g}"
        r[5].text = f"[{row['[0.025']:.3f}, {row['0.975]']:.3f}]"

    doc.add_paragraph(f"R-squared: {model.rsquared:.3f}; Adj. R-squared: {model.rsquared_adj:.3f}; F-statistic: {model.fvalue:.2f} (p={model.f_pvalue:.3g})")

    doc.add_heading("5. ANOVA (Type II)", level=2)
    anova_tbl = doc.add_table(rows=1, cols=len(anova_table.columns)+1)
    anova_tbl.rows[0].cells[0].text = "term"
    for j, col in enumerate(anova_table.columns):
        anova_tbl.rows[0].cells[j+1].text = col
    for i, idx in enumerate(anova_table.index):
        rc = anova_tbl.add_row().cells
        rc[0].text = str(idx)
        for j, col in enumerate(anova_table.columns):
            rc[j+1].text = f"{anova_table.iloc[i,j]:.4f}"

    doc.add_heading("6. Multicollinearity (VIF)", level=2)
    for _, row in vif_df.iterrows():
        doc.add_paragraph(f"{row['variable']}: VIF = {row['VIF']:.3f}")

    doc.add_heading("7. Diagnostic plots", level=2)
    doc.add_paragraph("Included: residuals vs fitted, QQ-plot of residuals, scatter matrix.")
    # attach images if present
    for img in (IMG_RESID, IMG_QQ, IMG_SCAT):
        if os.path.exists(img):
            try:
                doc.add_picture(img, width=Inches(6))
            except Exception:
                doc.add_paragraph(f"(Could not insert image {img} into DOCX; file exists.)")
        else:
            doc.add_paragraph(f"Image {img} not found.")

    doc.add_heading("8. Key findings and recommendations", level=2)
    doc.add_paragraph("• In this synthetic experiment, tillage depth, implement width, soil moisture, and speed show positive associations with fuel consumption (L/h).")
    doc.add_paragraph("• Interactions between depth & width and between soil moisture & speed capture combined effects and were included in the model.")
    doc.add_paragraph("• Diagnostics (QQ plot, residuals) should be checked in real data; heteroscedasticity can be present. Consider robust SEs if present in field data.")
    doc.add_paragraph("• For practical field reporting, consider fuel per area (L/ha), include tractor engine rpm, gear, soil texture, and implement setting in future experiments.")

    doc.save(DOCX_PATH)
    return DOCX_PATH

# -------------------------
# 5) Main flow
# -------------------------
def main():
    print("Generating synthetic dataset...")
    df = generate_synthetic_data()

    print(f"Saving CSV to {CSV_PATH} ...")
    df.to_csv(CSV_PATH, index=False)

    print("Running descriptive statistics...")
    desc = df.describe().T

    print("Computing correlation matrix...")
    corr = df.corr()

    print("Fitting regression model...")
    model, formula = fit_model(df)

    print("Computing ANOVA (Type II)...")
    anova_table = sm.stats.anova_lm(model, typ=2)

    print("Computing VIFs...")
    vif_df = compute_vif(df)

    print("Generating diagnostic plots...")
    make_plots(model, df)

    print("Building DOCX report...")
    docx_path = create_docx_report(df, desc, corr, model, formula, anova_table, vif_df)

    # Print short console summary
    print("\n=== Finished ===")
    print(f"CSV saved to: {CSV_PATH}")
    print(f"DOCX report saved to: {docx_path}")
    print(f"Residuals plot: {IMG_RESID}")
    print(f"QQ plot: {IMG_QQ}")
    print(f"Scatter matrix: {IMG_SCAT}")

    print("\nModel summary (brief):")
    print(model.summary().as_text())

if __name__ == "__main__":
    main()
