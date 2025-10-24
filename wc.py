import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import pearsonr
from docx import Document
from docx.shared import Inches

# === CONFIGURATION ===
np.random.seed(0)
N = 100  # sample size
IMG_DPI = 150  # smaller image file sizes

# === DATA GENERATION ===
mobile_usage = np.random.randint(0, 10, N)
gpa = 3.5 - 0.1 * mobile_usage + np.random.randn(N)

data = pd.DataFrame({
    'mobile_usage_hours': mobile_usage,
    'GPA': gpa
})

# === ANALYSIS ===
corr, p_value = pearsonr(data['mobile_usage_hours'], data['GPA'])

# === VISUALIZATION 1: Scatter Plot ===
plt.figure(figsize=(8, 5))
plt.scatter(data['mobile_usage_hours'], data['GPA'], alpha=0.7)
plt.xlabel('Mobile Usage (hours)')
plt.ylabel('GPA')
plt.title('Relationship between Mobile Usage and GPA')
plt.grid(alpha=0.3)
plt.tight_layout()
plt.savefig('relationship_plot.png', dpi=IMG_DPI)
plt.close()

# === VISUALIZATION 2: Average GPA by Usage Category ===
bins = [0, 2, 4, 6, np.inf]
labels = ['Low', 'Moderate', 'High', 'Very High']
data['usage_category'] = pd.cut(data['mobile_usage_hours'], bins=bins, labels=labels)

gpa_by_usage = data.groupby('usage_category', observed=True)['GPA'].mean().reset_index()

plt.figure(figsize=(6, 4))
plt.bar(gpa_by_usage['usage_category'], gpa_by_usage['GPA'], color='skyblue', edgecolor='black')
plt.xlabel('Mobile Usage Category')
plt.ylabel('Average GPA')
plt.title('Average GPA by Mobile Usage Category')
plt.grid(axis='y', alpha=0.3)
plt.tight_layout()
plt.savefig('gpa_by_usage.png', dpi=IMG_DPI)
plt.close()

# === EXPORT TO WORD DOCUMENT ===
doc = Document()
doc.add_heading('Mobile Device Usage and Academic Performance', level=0)

doc.add_paragraph(f"Correlation between mobile usage and GPA: {corr:.2f}")
doc.add_paragraph(f"p-value: {p_value:.4f}")

doc.add_heading('Relationship Plot', level=1)
doc.add_picture('relationship_plot.png', width=Inches(5.5))

doc.add_heading('Average GPA by Usage Category', level=1)
doc.add_picture('gpa_by_usage.png', width=Inches(5.5))

doc.save('Mobile_Usage_vs_GPA_Report.docx')

print("✅ Report generated: 'Mobile_Usage_vs_GPA_Report.docx'")
