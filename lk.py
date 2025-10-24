#!/usr/bin/env python3
"""
robust_alcohol_parenting_pipeline.py

Robust version of the synthetic data + analysis + reporting + optional dashboard script.
This version:
 - gracefully handles missing optional packages (dash/plotly),
 - ensures dummy columns exist,
 - avoids KeyError when categories are missing,
 - avoids crashes when embedding images in docx,
 - prints status messages for each major step.
"""

import os
import sys
import numpy as np
import pandas as pd
import matplotlib
# Use Agg backend for headless servers (safe for saving figures)
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# Optional visualization libs
try:
    import seaborn as sns
    SEABORN_AVAILABLE = True
except Exception:
    SEABORN_AVAILABLE = False

from scipy import stats
import statsmodels.api as sm
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression

# docx export
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Optional: dash & plotly
try:
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except Exception:
    PLOTLY_AVAILABLE = False

try:
    from dash import Dash, dcc, html  # type: ignore
    DASH_AVAILABLE = True
except Exception:
    DASH_AVAILABLE = False

# ---------------------------
# CONFIG & ASSUMPTIONS
# ---------------------------
RANDOM_SEED = 42
np.random.seed(RANDOM_SEED)
N = 5000  # sample size (adjustable)

ALCOHOL_CATEGORIES = ['None', 'Moderate', 'Heavy']
ALCOHOL_PROBS = [0.40, 0.45, 0.15]  # none, moderate, heavy

BASE_PARENTING_MEAN = 75
BASE_PARENTING_SD = 10
EFFECTS = {'None': 0.0, 'Moderate': -3.0, 'Heavy': -10.0}

RISK_MULTIPLIER = {'None': 1.0, 'Moderate': 1.5, 'Heavy': 3.0}

SES_MEAN = 0.0
SES_SD = 1.0

OUT_DIR = os.path.abspath(".")
CSV_NAME = os.path.join(OUT_DIR, "synthetic_alcohol_parenting.csv")
REPORT_NAME = os.path.join(OUT_DIR, "Alcohol_Parenting_Report.docx")

# plot filenames
VIOLIN_FNAME = os.path.join(OUT_DIR, "violin_parenting_by_alcohol.png")
SCATTER_FNAME = os.path.join(OUT_DIR, "scatter_parenting_risk.png")
BAR_FNAME = os.path.join(OUT_DIR, "bar_adverse_by_alcohol.png")

# ---------------------------
# 1. SYNTHESIZE DATA (robust)
# ---------------------------
print("1) Generating synthetic dataset...")

# sample categories
categories = np.random.choice(ALCOHOL_CATEGORIES, size=N, p=ALCOHOL_PROBS)

# Guarantee at least one of each category (prevents empty-group issues)
for cat in ALCOHOL_CATEGORIES:
    if cat not in categories:
        idx = np.random.randint(0, N)
        categories[idx] = cat

# build dataframe
df = pd.DataFrame({
    "alcohol_category": pd.Categorical(categories, categories=ALCOHOL_CATEGORIES),
    "ses_z": np.random.normal(SES_MEAN, SES_SD, N),
    "parent_age": np.random.normal(35, 6, N).clip(18, 80).round(1)
})

# parenting score
df["parenting_score"] = (
    np.random.normal(BASE_PARENTING_MEAN, BASE_PARENTING_SD, N)
    + df["alcohol_category"].map(EFFECTS).astype(float)
    + df["ses_z"] * 2.0
    + np.random.normal(0, 3, N)
)
df["parenting_score"] = df["parenting_score"].clip(0, 100)

# child risk continuous
base_risk = np.random.beta(a=2, b=14, size=N) * 0.4
df["child_risk_cont"] = base_risk * df["alcohol_category"].map(RISK_MULTIPLIER).astype(float)
df["child_risk_cont"] = df["child_risk_cont"] + (100 - df["parenting_score"]) * 0.002
df["child_risk_cont"] = df["child_risk_cont"].clip(0, 1)

# binary adverse
df["child_adverse"] = (np.random.rand(N) < df["child_risk_cont"]).astype(int)

# save csv
try:
    df.to_csv(CSV_NAME, index=False)
    print(f"   -> Saved CSV: {CSV_NAME}")
except Exception as e:
    print("   ! Failed to save CSV:", e)

# ---------------------------
# 2. DESCRIPTIVE STATISTICS
# ---------------------------
print("2) Descriptive statistics...")
desc = df.groupby("alcohol_category").agg(
    n=("alcohol_category", "size"),
    mean_parenting=("parenting_score", "mean"),
    sd_parenting=("parenting_score", "std"),
    adverse_rate=("child_adverse", "mean")
).reset_index()
print(desc.to_string(index=False))

# correlation
try:
    corr_res = stats.pearsonr(df["parenting_score"], df["child_risk_cont"])
    print(f"\n   Correlation parenting_score vs child_risk_cont: r={corr_res[0]:.3f}, p={corr_res[1]:.4f}")
except Exception as e:
    print("   ! Correlation failed:", e)

# ---------------------------
# 3. STATISTICAL TESTS & MODELS
# ---------------------------
print("3) Statistical tests & models...")

# ANOVA: require at least 2 groups with >1 observation
group_counts = df["alcohol_category"].value_counts()
valid_groups = [g for g in ALCOHOL_CATEGORIES if group_counts.get(g, 0) > 1]
if len(valid_groups) >= 2:
    try:
        samples = [df.loc[df["alcohol_category"] == g, "parenting_score"].values for g in valid_groups]
        anova_res = stats.f_oneway(*samples)
        print(f"   ANOVA on groups {valid_groups}: F={anova_res.statistic:.3f}, p={anova_res.pvalue:.4f}")
    except Exception as e:
        print("   ! ANOVA error:", e)
else:
    print("   ! Not enough groups with >1 observation to run ANOVA. Skipping.")

# safer dummy encoding for regression
df_lr = pd.get_dummies(df.copy(), columns=["alcohol_category"], drop_first=True)

# ensure expected dummy columns exist
expected_dummy_cols = [f"alcohol_category_{c}" for c in ALCOHOL_CATEGORIES if c != ALCOHOL_CATEGORIES[0]]
for col in expected_dummy_cols:
    if col not in df_lr.columns:
        df_lr[col] = 0

# Linear regression: parenting_score ~ alcohol dummies + ses + age
X_cols = expected_dummy_cols + ["ses_z", "parent_age"]
X = df_lr[X_cols]
X = sm.add_constant(X, has_constant='add')
y = df_lr["parenting_score"]
try:
    lm = sm.OLS(y, X).fit()
    print("\n   Linear regression results (coefficients):")
    print(lm.params.to_string())
except Exception as e:
    print("   ! Linear regression failed:", e)
    lm = None

# Logistic regression: child_adverse ~ features
X_log = df_lr[expected_dummy_cols + ["parenting_score", "ses_z", "parent_age"]]
y_log = df_lr["child_adverse"]
try:
    scaler = StandardScaler()
    X_log_scaled = scaler.fit_transform(X_log)
    logreg = LogisticRegression(max_iter=1000)
    logreg.fit(X_log_scaled, y_log)
    coefs = pd.DataFrame({"feature": X_log.columns, "coef": logreg.coef_.flatten()})
    print("\n   Logistic regression coefficients (log-odds scale):")
    print(coefs.to_string(index=False))
    df["pred_prob_adverse"] = logreg.predict_proba(scaler.transform(X_log))[:, 1]
except Exception as e:
    print("   ! Logistic regression failed:", e)
    df["pred_prob_adverse"] = np.nan

# ---------------------------
# 4. VISUALIZATIONS (saved)
# ---------------------------
print("4) Creating visualizations...")
if SEABORN_AVAILABLE:
    sns.set(style="whitegrid")
else:
    print("   - seaborn not available; falling back to matplotlib only")

# Violin / boxplot of parenting score by category
try:
    plt.figure(figsize=(8, 5))
    if SEABORN_AVAILABLE:
        sns.violinplot(x="alcohol_category", y="parenting_score", data=df, order=ALCOHOL_CATEGORIES)
    else:
        # simple boxplot fallback
        df.boxplot(column="parenting_score", by="alcohol_category", grid=False)
        plt.suptitle("")
    plt.title("Parenting Quality by Parental Alcohol Use")
    plt.tight_layout()
    plt.savefig(VIOLIN_FNAME, dpi=150)
    plt.close()
    print(f"   -> Saved: {VIOLIN_FNAME}")
except Exception as e:
    print("   ! Failed to create violin plot:", e)

# Scatter parenting vs child risk with smoothing
try:
    plt.figure(figsize=(7, 5))
    if SEABORN_AVAILABLE:
        sns.scatterplot(x="parenting_score", y="child_risk_cont", hue="alcohol_category", data=df, alpha=0.4)
        sns.regplot(x="parenting_score", y="child_risk_cont", data=df, scatter=False, lowess=True, color="k")
    else:
        plt.scatter(df["parenting_score"], df["child_risk_cont"], alpha=0.4)
    plt.title("Child Risk vs Parenting Score")
    plt.tight_layout()
    plt.savefig(SCATTER_FNAME, dpi=150)
    plt.close()
    print(f"   -> Saved: {SCATTER_FNAME}")
except Exception as e:
    print("   ! Failed to create scatter plot:", e)

# Bar chart adverse rate by alcohol category
try:
    order = ALCOHOL_CATEGORIES
    rates = df.groupby("alcohol_category")["child_adverse"].mean().reindex(order)
    plt.figure(figsize=(6, 4))
    if SEABORN_AVAILABLE:
        sns.barplot(x=rates.index, y=rates.values)
    else:
        plt.bar(rates.index, rates.values)
    plt.ylabel("Adverse Outcome Rate")
    plt.title("Adverse Rate by Alcohol Category")
    plt.tight_layout()
    plt.savefig(BAR_FNAME, dpi=150)
    plt.close()
    print(f"   -> Saved: {BAR_FNAME}")
except Exception as e:
    print("   ! Failed to create bar plot:", e)

# ---------------------------
# 5. EXPORT REPORT (docx) - optional
# ---------------------------
print("5) Exporting report (docx)...")
if not DOCX_AVAILABLE:
    print("   - python-docx not installed. Skipping report export.")
else:
    try:
        doc = Document()
        doc.add_heading("Effect of Parental Alcohol Use on Parenting (Synthetic Study)", level=0)
        doc.add_paragraph(f"This synthetic dataset has {len(df):,} simulated parent observations.")
        doc.add_heading("Key descriptives", level=1)
        for _, row in desc.iterrows():
            doc.add_paragraph(
                f"{row['alcohol_category']}: n={int(row['n'])}, mean parenting={row['mean_parenting']:.2f}, "
                f"SD={row['sd_parenting']:.2f}, adverse_rate={row['adverse_rate']:.3f}"
            )

        doc.add_heading("Statistical results", level=1)
        if 'anova_res' in locals():
            try:
                doc.add_paragraph(f"ANOVA (parenting_score by alcohol_category): F={anova_res.statistic:.3f}, p={anova_res.pvalue:.4f}")
            except Exception:
                doc.add_paragraph("ANOVA was not run due to insufficient group sizes.")
        if lm is not None:
            doc.add_paragraph("Linear regression coefficients (parenting_score ~ alcohol + ses + age):")
            for term, coef in lm.params.items():
                doc.add_paragraph(f"  {term}: {coef:.3f}")
        if 'coefs' in locals():
            doc.add_paragraph("Logistic regression coefficients (child_adverse):")
            for _, r in coefs.iterrows():
                doc.add_paragraph(f"  {r['feature']}: {r['coef']:.3f}")

        doc.add_heading("Figures", level=1)
        for fname, caption in [(VIOLIN_FNAME, "Parenting quality distribution by alcohol use"),
                               (SCATTER_FNAME, "Parenting score vs child risk"),
                               (BAR_FNAME, "Adverse outcome rates by alcohol category")]:
            if os.path.exists(fname):
                doc.add_paragraph(caption)
                # scale images reasonably
                try:
                    doc.add_picture(fname, width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"(Could not embed {fname}: {e})")
            else:
                doc.add_paragraph(f"(Missing figure: {fname})")

        doc.save(REPORT_NAME)
        print(f"   -> Saved report: {REPORT_NAME}")
    except Exception as e:
        print("   ! Failed to create Word report:", e)

# ---------------------------
# 6. DASHBOARD (optional)
# ---------------------------
if DASH_AVAILABLE and PLOTLY_AVAILABLE:
    print("6) Dash & Plotly available — launching dashboard (http://127.0.0.1:8050/)...")
    try:
        app = Dash(__name__)
        fig_violin = px.violin(df, x="alcohol_category", y="parenting_score", box=True, points="all",
                               category_orders={"alcohol_category": ALCOHOL_CATEGORIES})
        # sample for scatter if large
        sample_df = df.sample(min(len(df), 2000), random_state=RANDOM_SEED)
        fig_scatter = px.scatter(sample_df, x="parenting_score", y="child_risk_cont", color="alcohol_category",
                                 title="Parenting vs Child Risk (sample)")
        fig_bar = px.bar(x=ALCOHOL_CATEGORIES, y=df.groupby("alcohol_category")["child_adverse"].mean().reindex(ALCOHOL_CATEGORIES).values,
                         labels={"x": "Alcohol category", "y": "Adverse rate"}, title="Adverse rate by alcohol category")

        app.layout = html.Div([
            html.H1("Alcohol & Parenting — Synthetic Study Dashboard"),
            dcc.Tabs([
                dcc.Tab(label="Overview", children=[
                    dcc.Graph(figure=fig_violin),
                    dcc.Graph(figure=fig_bar)
                ]),
                dcc.Tab(label="Relationships", children=[
                    dcc.Graph(figure=fig_scatter),
                    html.Pre(lm.summary().as_text() if lm is not None else "No regression summary available",
                             style={"whiteSpace": "pre-wrap", "overflowX": "auto"})
                ])
            ])
        ])

        if __name__ == "__main__":
            # run server
            app.run_server(debug=False)
    except Exception as e:
        print("   ! Failed to start dashboard:", e)
else:
    print("6) Dash/Plotly not available or skipped — dashboard not started.")
    if not DASH_AVAILABLE:
        print("   - dash missing; install with: pip install dash")
    if not PLOTLY_AVAILABLE:
        print("   - plotly missing; install with: pip install plotly")

print("\nAll done. Check the CSV, PNGs and (if available) the Word report in the current directory.")
