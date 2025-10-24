"""
alcohol_parenting_pipeline.py

Generates synthetic data informed by published studies on parental alcohol use,
analyzes effects on parenting quality and child-risk outcomes, visualizes results,
serves a simple Plotly Dash dashboard, and exports a Word report with figures.

Cited sources used to parameterize the simulation (examples):
- Parental alcohol use reviews and epidemiology (cross-sectional & longitudinal). 1
- Studies linking parental drinking to poorer parenting and adverse child outcomes. 2
- Recent meta-analysis / association to maltreatment. 3

Note: This is a simulation synthesized from reported associations and prevalences,
not raw participant-level data from the cited papers (which often are not available as CSV).
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
import statsmodels.api as sm
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from docx import Document
from docx.shared import Inches
import plotly.express as px
import plotly.graph_objects as go
from dash import Dash, dcc, html, Input, Output
import os

# ---------------------------
# 1. CONFIG & ASSUMPTIONS
# ---------------------------
RANDOM_SEED = 42
np.random.seed(RANDOM_SEED)
N = 5000  # simulated families / parent observations

# Prevalence assumptions (informed by cited literature):
# - Many parents drink; a minority have heavy/problem drinking (~4-8% with AUD per some reports).
# We'll model three categories: 'None', 'Moderate', 'Heavy'
ALCOHOL_PROBS = [0.40, 0.45, 0.15]  # none, moderate, heavy

# Parenting quality baseline (0-100): Higher is better.
BASE_PARENTING_MEAN = 75
BASE_PARENTING_SD = 10

# Effects: decrement in parenting quality by alcohol category (approx. small-medium effect)
# informed by literature that heavy drinking reduces warmth, monitoring and increases inconsistent discipline.
EFFECTS = {
    'None': 0.0,
    'Moderate': -3.0,   # moderate drinking reduces parenting score by ~3 points on average
    'Heavy': -10.0      # heavy / AUD reduces parenting score by ~10 points on average
}

# Child adverse outcome risk - we'll synthesize a continuous risk score and a binary "adverse" flag.
# Heavy parental drinking increases risk substantially (ORs vary in literature; we'll set plausible effects).
BASE_CHILD_RISK_MEAN = 0.12  # baseline probability of an adverse outcome ~12%
RISK_MULTIPLIER = {'None': 1.0, 'Moderate': 1.5, 'Heavy': 3.0}

# Socioeconomic status as a confounder
SES_MEAN = 0.0
SES_SD = 1.0

# ---------------------------
# 2. SYNTHESIZE DATA
# ---------------------------
# sample alcohol category
categories = np.random.choice(['None', 'Moderate', 'Heavy'], size=N, p=ALCOHOL_PROBS)

# socioeconomic status (standardized)
ses = np.random.normal(SES_MEAN, SES_SD, N)

# parental age
parent_age = np.random.normal(35, 6, N).clip(18, 80).round(1)

# parenting quality score (0-100)
parenting_score = (
    np.random.normal(BASE_PARENTING_MEAN, BASE_PARENTING_SD, N)
    + np.array([EFFECTS[c] for c in categories])
    + (ses * 2.0)  # assume higher SES -> slightly higher parenting score
    + np.random.normal(0, 3, N)  # additional noise
)
parenting_score = parenting_score.clip(0, 100)

# child risk as continuous probability-like (0-1)
base_risk = np.random.beta(a=2, b=14, size=N) * 0.4  # shape to produce low baseline probabilities
child_risk_cont = base_risk * np.array([RISK_MULTIPLIER[c] for c in categories])

# Let poorer parenting also increase risk (partial mediation)
# combine multiplicative and additive effects
child_risk_cont = child_risk_cont + (100 - parenting_score) * 0.002  # each 1-point drop adds 0.002 to risk
child_risk_cont = np.clip(child_risk_cont, 0, 1)

# Binary adverse outcome (e.g., serious child-related adverse indicator)
child_adverse = np.array([np.random.rand() < p for p in child_risk_cont]).astype(int)

# Assemble into DataFrame
df = pd.DataFrame({
    'alcohol_category': categories,
    'ses_z': ses,
    'parent_age': parent_age,
    'parenting_score': parenting_score,
    'child_risk_cont': child_risk_cont,
    'child_adverse': child_adverse
})

# Save synthetic CSV
df.to_csv('synthetic_alcohol_parenting.csv', index=False)
print("Saved: synthetic_alcohol_parenting.csv (N=%d)" % N)

# ---------------------------
# 3. DESCRIPTIVE STATISTICS
# ---------------------------
desc = df.groupby('alcohol_category').agg(
    n=('alcohol_category', 'size'),
    mean_parenting=('parenting_score', 'mean'),
    sd_parenting=('parenting_score', 'std'),
    adverse_rate=('child_adverse', 'mean')
).reset_index()
print("\nDescriptives by alcohol category:\n", desc)

# Overall correlations
corr_parenting_risk = stats.pearsonr(df['parenting_score'], df['child_risk_cont'])
print("\nCorrelation parenting_score vs child_risk_cont: r=%.3f p=%.4f" % corr_parenting_risk)

# ---------------------------
# 4. STATISTICAL TESTS & MODELS
# ---------------------------
# ANOVA: parenting_score by alcohol_category
anova_res = stats.f_oneway(
    df.loc[df['alcohol_category'] == 'None', 'parenting_score'],
    df.loc[df['alcohol_category'] == 'Moderate', 'parenting_score'],
    df.loc[df['alcohol_category'] == 'Heavy', 'parenting_score']
)
print("\nANOVA parenting_score by alcohol_category: F=%.3f p=%.4f" % (anova_res.statistic, anova_res.pvalue))

# Linear regression: parenting_score ~ alcohol_category + ses + parent_age
# encode categories
df_lr = pd.get_dummies(df, columns=['alcohol_category'], drop_first=True)  # drop 'None' as reference
X = df_lr[['alcohol_category_Moderate', 'alcohol_category_Heavy', 'ses_z', 'parent_age']]
X = sm.add_constant(X)
y = df_lr['parenting_score']
lm = sm.OLS(y, X).fit()
print("\nLinear regression summary (parenting_score ~ alcohol + ses + age):\n")
print(lm.summary().tables[1])

# Logistic regression: child_adverse ~ alcohol_category + parenting_score + ses + parent_age
from sklearn.linear_model import LogisticRegression
X_log = df_lr[['alcohol_category_Moderate', 'alcohol_category_Heavy', 'parenting_score', 'ses_z', 'parent_age']]
y_log = df_lr['child_adverse']
scaler = StandardScaler()
X_log_scaled = scaler.fit_transform(X_log)
logreg = LogisticRegression(max_iter=1000)
logreg.fit(X_log_scaled, y_log)
# extract odds-like coefficients (approx)
coefs = pd.DataFrame({
    'feature': X_log.columns,
    'coef': logreg.coef_.flatten()
})
print("\nLogistic regression (child_adverse) coefficients (log-odds scale):\n", coefs)

# compute predicted probabilities for plotting
df['pred_prob_adverse'] = logreg.predict_proba(scaler.transform(X_log))[:,1]

# ---------------------------
# 5. VISUALIZATIONS (static saves)
# ---------------------------
sns.set(style="whitegrid")

# Violin of parenting score by alcohol category
plt.figure(figsize=(8,5))
sns.violinplot(x='alcohol_category', y='parenting_score', data=df, order=['None','Moderate','Heavy'])
plt.title('Parenting Quality by Parental Alcohol Use')
plt.xlabel('Alcohol category')
plt.ylabel('Parenting quality score (0-100)')
plt.tight_layout()
plt.savefig('violin_parenting_by_alcohol.png', dpi=150)
plt.close()

# Scatter parenting vs risk with jitter for clarity & smoothing trend
plt.figure(figsize=(7,5))
sns.scatterplot(x='parenting_score', y='child_risk_cont', hue='alcohol_category', data=df, alpha=0.4)
sns.regplot(x='parenting_score', y='child_risk_cont', data=df, scatter=False, lowess=True, color='k')
plt.title('Child risk (continuous) vs Parenting score')
plt.xlabel('Parenting score')
plt.ylabel('Child risk (0-1)')
plt.tight_layout()
plt.savefig('scatter_parenting_risk.png', dpi=150)
plt.close()

# Bar chart: adverse rate by alcohol category
plt.figure(figsize=(6,4))
order = ['None','Moderate','Heavy']
rates = df.groupby('alcohol_category')['child_adverse'].mean().reindex(order)
sns.barplot(x=rates.index, y=rates.values)
plt.ylabel('Adverse outcome rate')
plt.title('Adverse outcome rate by parental alcohol category')
plt.ylim(0, rates.max()*1.4)
plt.tight_layout()
plt.savefig('bar_adverse_by_alcohol.png', dpi=150)
plt.close()

print("Saved plots: violin_parenting_by_alcohol.png, scatter_parenting_risk.png, bar_adverse_by_alcohol.png")

# ---------------------------
# 6. EXPORT REPORT (docx)
# ---------------------------
doc = Document()
doc.add_heading('Effect of Parental Alcohol Use on Parenting (Synthetic study)', level=0)
doc.add_paragraph('This report uses a synthetic dataset parameterized using published literature that reports associations between parental drinking and parenting or child outcomes. The dataset contains %d simulated parent observations.' % N)

doc.add_heading('Key descriptives', level=1)
for _, row in desc.iterrows():
    doc.add_paragraph(f"{row['alcohol_category']}: n={int(row['n'])}, mean parenting={row['mean_parenting']:.2f}, SD={row['sd_parenting']:.2f}, adverse_rate={row['adverse_rate']:.3f}")

doc.add_heading('Statistical results', level=1)
doc.add_paragraph('ANOVA for parenting_score by alcohol_category: F=%.3f, p=%.4f' % (anova_res.statistic, anova_res.pvalue))
doc.add_paragraph('Linear regression (parenting_score ~ alcohol + ses + age):')
# add top coefficients from lm
lm_coefs = lm.params.reset_index()
lm_coefs.columns = ['term', 'coef']
for _, r in lm_coefs.iterrows():
    doc.add_paragraph(f"  {r['term']}: {r['coef']:.3f}")

doc.add_paragraph('Logistic regression (child_adverse) coefficients (log-odds scale):')
for _, r in coefs.iterrows():
    doc.add_paragraph(f"  {r['feature']}: {r['coef']:.3f}")

# add figures
doc.add_heading('Figures', level=1)
doc.add_paragraph('Parenting quality distribution by alcohol use:')
doc.add_picture('violin_parenting_by_alcohol.png', width=Inches(6))
doc.add_page_break()
doc.add_paragraph('Relationship between parenting quality and child risk:')
doc.add_picture('scatter_parenting_risk.png', width=Inches(6))
doc.add_paragraph('Adverse outcome rates by parental alcohol category:')
doc.add_picture('bar_adverse_by_alcohol.png', width=Inches(6))

report_name = 'Alcohol_Parenting_Report.docx'
doc.save(report_name)
print(f"Saved report: {report_name}")

# ---------------------------
# 7. DASHBOARD (Plotly Dash)
# ---------------------------
# Build a small Dash app to explore the synthesized data.
app = Dash(__name__)
server = app.server

fig_violin = px.violin(df, x='alcohol_category', y='parenting_score', box=True, points='all',
                       category_orders={'alcohol_category':['None','Moderate','Heavy']},
                       labels={'alcohol_category':'Alcohol category', 'parenting_score':'Parenting score'})
fig_scatter = px.scatter(df.sample(2000), x='parenting_score', y='child_risk_cont', color='alcohol_category',
                         labels={'child_risk_cont':'Child risk (continuous)'}, title='Parenting vs Child Risk (sample)')
fig_bar = px.bar(x=order, y=rates.values, labels={'x':'Alcohol category', 'y':'Adverse rate'}, title='Adverse rate by alcohol category')

app.layout = html.Div([
    html.H1("Alcohol & Parenting — Synthetic Study Dashboard"),
    dcc.Tabs([
        dcc.Tab(label='Overview', children=[
            html.Div([
                dcc.Graph(figure=fig_violin),
                dcc.Graph(figure=fig_bar)
            ])
        ]),
        dcc.Tab(label='Relationships', children=[
            html.Div([
                dcc.Graph(figure=fig_scatter),
                html.H4('Regression summary (parenting_score ~ alcohol + ses + age)'),
                html.Pre(lm.summary().as_text(), style={'whiteSpace': 'pre-wrap', 'overflowX': 'auto'})
            ])
        ]),
        dcc.Tab(label='Data & Export', children=[
            html.H4('Download / Export'),
            html.P('The script has already saved a Word report and the CSV in the working directory.'),
            html.Ul([
                html.Li('synthetic_alcohol_parenting.csv'),
                html.Li('violin_parenting_by_alcohol.png'),
                html.Li('scatter_parenting_risk.png'),
                html.Li('bar_adverse_by_alcohol.png'),
                html.Li(report_name)
            ]),
            html.P('To re-generate and download modified reports, re-run the script with different parameters.')
        ])
    ])
])

if __name__ == '__main__':
    print("Launching dashboard at http://127.0.0.1:8050/ . Press Ctrl+C to stop.")
    app.run_server(debug=False)
