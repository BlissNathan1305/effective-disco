import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import pearsonr
from docx import Document
from docx.shared import Inches

# Generate a random dataset
np.random.seed(0)
mobile_usage_hours = np.random.randint(0, 10, 100)
GPA = 3.5 - 0.1 * mobile_usage_hours + np.random.randn(100)

data = pd.DataFrame({
    'mobile_usage_hours': mobile_usage_hours,
    'GPA': GPA
})

# Calculate correlation between mobile usage and GPA
correlation, p_value = pearsonr(data['mobile_usage_hours'], data['GPA'])

# Create visualizations
plt.figure(figsize=(10, 6))
plt.scatter(data['mobile_usage_hours'], data['GPA'])
plt.xlabel('Mobile Usage (hours)')
plt.ylabel('GPA')
plt.title('Relationship between Mobile Usage and GPA')
plt.savefig('relationship_plot.png', bbox_inches='tight')

data['usage_category'] = pd.cut(data['mobile_usage_hours'], bins=[0, 2, 4, 6, np.inf], labels=['Low', 'Moderate', 'High', 'Very High'])
gpa_by_usage = data.groupby('usage_category')['GPA'].mean().reset_index()

plt.figure(figsize=(8, 6))
plt.bar(gpa_by_usage['usage_category'], gpa_by_usage['GPA'])
plt.xlabel('Mobile Usage Category')
plt.ylabel('Average GPA')
plt.title('Average GPA by Mobile Usage Category')
plt.savefig('gpa_by_usage.png', bbox_inches='tight')

# Export findings to Word document
document = Document()
document.add_heading('Mobile Device Usage and Academic Performance', 0)

document.add_paragraph(f"Correlation between mobile usage and GPA: {correlation:.2f}")
document.add_paragraph(f"p-value: {p_value:.2f}")

document.add_heading('Relationship Plot', level=1)
document
